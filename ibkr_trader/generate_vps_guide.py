"""
Generates VPS Deployment Guide for IBKR Auto-Trader as a Word document.
Run: python generate_vps_guide.py
Output: ibkr_vps_deployment_guide.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=16, bold=True, color=(31, 73, 125))
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    set_font(run, size=13, bold=True, color=(0, 112, 192))
    return p

def body(text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(1.2 + level * 0.8)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def code_block(lines):
    """Grey-background paragraph for each line of code."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.8)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F2F2F2')
        pPr.append(shd)
        run = p.add_run(line if line else " ")
        set_font(run, name="Courier New", size=9, color=(50, 50, 50))

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFF4CE')
    pPr.append(shd)
    run = p.add_run("NOTE: " + text)
    set_font(run, size=10, color=(124, 83, 0))
    run.font.bold = True

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                set_font(run, size=10, bold=True, color=(255, 255, 255))
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F497D')
        tcPr.append(shd)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        fill = 'DEEAF1' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = val
            for para in cell.paragraphs:
                for run in para.runs:
                    set_font(run, size=10)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)
    doc.add_paragraph()  # spacing after table

# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run("IBKR Options Auto-Trader")
set_font(run, size=28, bold=True, color=(31, 73, 125))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("VPS Deployment Guide")
set_font(run, size=20, bold=True, color=(0, 112, 192))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
run = p.add_run("FastAPI + React + IB Gateway + IBC + Systemd")
set_font(run, size=12, color=(89, 89, 89))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
run = p.add_run("v6 — June 2026")
set_font(run, size=11, color=(128, 128, 128))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════
heading1("Overview")
body(
    "This guide covers end-to-end deployment of the IBKR Options Auto-Trader on a Linux VPS. "
    "The system includes a FastAPI backend, single-page React frontend, IB Gateway (headless via IBC), "
    "SQLite trade journal, XGBoost learning engine, and a watchdog process for 24/7 reliability."
)

body("Architecture on VPS:")
bullet("Nginx (port 80/443) — serves frontend HTML, proxies /api/* to FastAPI")
bullet("FastAPI backend (port 8000, internal) — scanning, order placement, auto-trader loop")
bullet("IB Gateway (port 4002 paper / 4001 live, localhost only) — IBKR API bridge")
bullet("IBC — automates IB Gateway login on boot, no manual intervention needed")
bullet("Xvfb — virtual display required by IB Gateway (Java GUI)")
bullet("Watchdog — health-checks backend every 60s, auto-restarts on failure")
bullet("Systemd — manages all services, starts them on reboot")

# ═══════════════════════════════════════════════════════════════════════════════
# VPS REQUIREMENTS
# ═══════════════════════════════════════════════════════════════════════════════
heading1("Step 0 — VPS Requirements")

add_table(
    ["Resource", "Minimum", "Recommended"],
    [
        ["OS",       "Ubuntu 22.04 LTS",  "Ubuntu 22.04 LTS"],
        ["RAM",      "4 GB",              "8 GB"],
        ["CPU",      "2 vCPU",            "4 vCPU"],
        ["Disk",     "20 GB SSD",         "40 GB SSD"],
        ["Provider", "Any",               "DigitalOcean / Linode / AWS EC2"],
        ["Est. Cost","~$12/month",         "~$24/month"],
    ]
)

add_table(
    ["Port", "Service", "Exposed to Internet?"],
    [
        ["22",       "SSH",                        "Yes — key auth only"],
        ["80 / 443", "Nginx (frontend + API)",     "Yes"],
        ["8000",     "FastAPI (direct)",            "No — localhost only"],
        ["4002",     "IB Gateway paper trading",   "No — localhost only"],
        ["4001",     "IB Gateway live trading",    "No — localhost only"],
    ]
)

# ═══════════════════════════════════════════════════════════════════════════════
# STEP 1
# ═══════════════════════════════════════════════════════════════════════════════
heading1("Step 1 — Provision & Secure the VPS")
body("After creating the VPS with Ubuntu 22.04, SSH in as root and run:")

code_block([
    "# Create non-root user",
    "adduser trader",
    "usermod -aG sudo trader",
    "su - trader",
    "",
    "# Firewall — allow only SSH and web traffic",
    "sudo ufw allow OpenSSH",
    "sudo ufw allow 80",
    "sudo ufw allow 443",
    "sudo ufw enable",
    "",
    "# Disable password SSH login (use keys instead)",
    "sudo nano /etc/ssh/sshd_config",
    "# Set: PasswordAuthentication no",
    "sudo systemctl restart sshd",
])

note("Generate an SSH key pair locally (ssh-keygen) and copy the public key to the VPS "
     "before disabling password auth: ssh-copy-id trader@YOUR_VPS_IP")

# ═══════════════════════════════════════════════════════════════════════════════
# STEP 2
# ═══════════════════════════════════════════════════════════════════════════════
heading1("Step 2 — Install System Dependencies")

code_block([
    "sudo apt update && sudo apt upgrade -y",
    "",
    "# Python 3.11, Java 17 (IB Gateway), Nginx, virtual display tools",
    "sudo apt install -y \\",
    "  python3.11 python3.11-venv python3-pip \\",
    "  openjdk-17-jre \\",
    "  nginx \\",
    "  xvfb x11vnc \\",
    "  unzip curl git wget",
    "",
    "# Verify Java",
    "java -version",
    "",
    "# Verify Python",
    "python3.11 --version",
])

# ═══════════════════════════════════════════════════════════════════════════════
# STEP 3
# ═══════════════════════════════════════════════════════════════════════════════
heading1("Step 3 — Install IB Gateway + IBC")

heading2("3a. Download and Install IB Gateway")
body("IB Gateway is the headless-friendly version of TWS — it uses less RAM and is designed for API trading.")

code_block([
    "mkdir -p ~/ibkr && cd ~/ibkr",
    "",
    "# Download IB Gateway stable installer",
    "wget 'https://download2.interactivebrokers.com/installers/ibgateway/stable-standalone/ibgateway-stable-standalone-linux-x64.sh'",
    "chmod +x ibgateway-stable-standalone-linux-x64.sh",
    "",
    "# Start a virtual display (installer needs a GUI)",
    "Xvfb :99 -screen 0 1024x768x16 &",
    "export DISPLAY=:99",
    "",
    "# Run installer — accept defaults, install to ~/Jts",
    "bash ibgateway-stable-standalone-linux-x64.sh",
])

heading2("3b. Download and Configure IBC")
body("IBC automates the IB Gateway login so it can start unattended after a reboot.")

code_block([
    "cd ~/ibkr",
    "wget https://github.com/IbcAlpha/IBC/releases/latest/download/IBC.zip",
    "unzip IBC.zip -d ibc",
    "chmod +x ~/ibkr/ibc/scripts/*.sh",
    "",
    "# Edit the IBC config file",
    "nano ~/ibkr/ibc/config.ini",
])

body("Key settings to configure in config.ini:")
code_block([
    "IbLoginId=YOUR_IBKR_USERNAME",
    "IbPassword=YOUR_IBKR_PASSWORD",
    "TradingMode=paper              # change to 'live' when ready",
    "AcceptEulaMode=accept",
    "ReadOnlyLogin=no",
    "AcceptIncomingConnectionAction=accept",
    "ExistingSessionDetectedAction=primary",
    "ReloginAfterSecondFactorAuthenticationTimeout=yes",
])

note("IB Gateway uses ports 4002 (paper) and 4001 (live) by default — "
     "different from TWS which uses 7497/7496. Update the backend .env accordingly.")

heading2("3c. Configure IB Gateway API Settings")
body("In IB Gateway's settings, ensure these are set (first manual launch to configure):")
bullet("Enable ActiveX and Socket Clients: YES")
bullet("Socket port: 4002 (paper) or 4001 (live)")
bullet("Allow connections from localhost only: YES")
bullet("Read-Only API: NO")

# ═══════════════════════════════════════════════════════════════════════════════
# STEP 4
# ═══════════════════════════════════════════════════════════════════════════════
heading1("Step 4 — Deploy Your Code")

heading2("4a. Package and Upload from Windows")
body("Run this on your local Windows machine (PowerShell):")

code_block([
    "# Navigate to project",
    "cd C:\\Projects\\GenAI-Projects\\ibkr_trader",
    "",
    "# Create archive (exclude venv and cache)",
    "tar -czf ibkr_trader.tar.gz `",
    "  --exclude=venv --exclude=__pycache__ `",
    "  --exclude='*.pyc' --exclude='*.db' .",
    "",
    "# Upload to VPS",
    "scp ibkr_trader.tar.gz trader@YOUR_VPS_IP:~/",
])

heading2("4b. Extract and Set Up on VPS")
code_block([
    "mkdir -p ~/ibkr_trader && cd ~/ibkr_trader",
    "tar -xzf ~/ibkr_trader.tar.gz",
    "",
    "# Create Python virtual environment",
    "python3.11 -m venv venv",
    "source venv/bin/activate",
    "",
    "# Install all Python dependencies",
    "pip install --upgrade pip",
    "pip install fastapi uvicorn ib_insync yfinance pandas numpy \\",
    "            xgboost scikit-learn scipy aiohttp python-multipart",
])

# ═══════════════════════════════════════════════════════════════════════════════
# STEP 5
# ═══════════════════════════════════════════════════════════════════════════════
heading1("Step 5 — Environment Configuration")

heading2("5a. Create the .env file (NEVER commit this to git)")
code_block([
    "nano ~/ibkr_trader/backend/.env",
])

body("Paste the following content, replacing values as appropriate:")
code_block([
    "TWS_HOST=127.0.0.1",
    "TWS_PORT=4002          # paper: 4002 | live: 4001",
    "TWS_CLIENT_ID=1",
])

heading2("5b. Update main.py to read from environment")
body("At the top of backend/main.py, replace the hardcoded constants with:")
code_block([
    "import os",
    "TWS_HOST      = os.getenv('TWS_HOST', '127.0.0.1')",
    "TWS_PORT      = int(os.getenv('TWS_PORT', '4002'))",
    "TWS_CLIENT_ID = int(os.getenv('TWS_CLIENT_ID', '1'))",
])

heading2("5c. Update frontend API base URL")
body("In frontend/index.html, find the apiFetch function and update the base URL "
     "so the browser calls Nginx's /api prefix instead of localhost:8000 directly:")
code_block([
    "# Find this line in index.html:",
    "const BASE = 'http://localhost:8000';",
    "",
    "# Replace with:",
    "const BASE = '/api';",
])

# ═══════════════════════════════════════════════════════════════════════════════
# STEP 6
# ═══════════════════════════════════════════════════════════════════════════════
heading1("Step 6 — Configure Nginx")

code_block([
    "sudo nano /etc/nginx/sites-available/ibkr",
])

body("Paste this configuration:")
code_block([
    "server {",
    "    listen 80;",
    "    server_name YOUR_VPS_IP;  # replace with domain if available",
    "",
    "    # Serve frontend",
    "    location / {",
    "        root /home/trader/ibkr_trader/frontend;",
    "        index index.html;",
    "        try_files $uri $uri/ /index.html;",
    "    }",
    "",
    "    # Proxy API calls to FastAPI",
    "    location /api/ {",
    "        rewrite ^/api(/.*)$ $1 break;",
    "        proxy_pass http://127.0.0.1:8000;",
    "        proxy_set_header Host $host;",
    "        proxy_set_header X-Real-IP $remote_addr;",
    "        proxy_read_timeout 120s;",
    "    }",
    "",
    "    # WebSocket support for streaming bars",
    "    location /ws {",
    "        proxy_pass http://127.0.0.1:8000;",
    "        proxy_http_version 1.1;",
    "        proxy_set_header Upgrade $http_upgrade;",
    "        proxy_set_header Connection upgrade;",
    "    }",
    "}",
])

code_block([
    "# Enable site and test config",
    "sudo ln -s /etc/nginx/sites-available/ibkr /etc/nginx/sites-enabled/",
    "sudo nginx -t",
    "sudo systemctl reload nginx",
])

# ═══════════════════════════════════════════════════════════════════════════════
# STEP 7
# ═══════════════════════════════════════════════════════════════════════════════
heading1("Step 7 — Create Systemd Services")
body("Systemd ensures all processes start automatically on server reboot.")

heading2("7a. IB Gateway Service")
code_block([
    "sudo nano /etc/systemd/system/ibgateway.service",
])
code_block([
    "[Unit]",
    "Description=IB Gateway (headless via IBC)",
    "After=network.target",
    "",
    "[Service]",
    "User=trader",
    "Environment=DISPLAY=:99",
    "ExecStartPre=/usr/bin/Xvfb :99 -screen 0 1024x768x16",
    "ExecStart=/home/trader/ibkr/ibc/scripts/ibcstart.sh",
    "Restart=always",
    "RestartSec=30",
    "TimeoutStopSec=30",
    "",
    "[Install]",
    "WantedBy=multi-user.target",
])

heading2("7b. Backend Service")
code_block([
    "sudo nano /etc/systemd/system/ibkr-backend.service",
])
code_block([
    "[Unit]",
    "Description=IBKR Trading Backend (FastAPI)",
    "After=network.target ibgateway.service",
    "Requires=ibgateway.service",
    "",
    "[Service]",
    "User=trader",
    "WorkingDirectory=/home/trader/ibkr_trader/backend",
    "EnvironmentFile=/home/trader/ibkr_trader/backend/.env",
    "ExecStart=/home/trader/ibkr_trader/venv/bin/python \\",
    "          -m uvicorn main:app --host 0.0.0.0 --port 8000",
    "Restart=always",
    "RestartSec=10",
    "TimeoutStartSec=60",
    "",
    "[Install]",
    "WantedBy=multi-user.target",
])

heading2("7c. Watchdog Service")
code_block([
    "sudo nano /etc/systemd/system/ibkr-watchdog.service",
])
code_block([
    "[Unit]",
    "Description=IBKR Backend Watchdog",
    "After=ibkr-backend.service",
    "Requires=ibkr-backend.service",
    "",
    "[Service]",
    "User=trader",
    "WorkingDirectory=/home/trader/ibkr_trader",
    "ExecStart=/home/trader/ibkr_trader/venv/bin/python watchdog.py",
    "Restart=always",
    "RestartSec=15",
    "",
    "[Install]",
    "WantedBy=multi-user.target",
])

heading2("7d. Enable and Start All Services")
code_block([
    "sudo systemctl daemon-reload",
    "",
    "# Enable services to start on boot",
    "sudo systemctl enable ibgateway ibkr-backend ibkr-watchdog",
    "",
    "# Start in order (gateway first, wait 30s for login, then backend)",
    "sudo systemctl start ibgateway",
    "sleep 30",
    "sudo systemctl start ibkr-backend",
    "sleep 10",
    "sudo systemctl start ibkr-watchdog",
])

# ═══════════════════════════════════════════════════════════════════════════════
# STEP 8
# ═══════════════════════════════════════════════════════════════════════════════
heading1("Step 8 — Verify the Deployment")

code_block([
    "# Check all services are running",
    "sudo systemctl status ibgateway",
    "sudo systemctl status ibkr-backend",
    "sudo systemctl status ibkr-watchdog",
    "",
    "# Check backend API is responding",
    "curl http://localhost:8000/health",
    "curl http://localhost:8000/status",
    "",
    "# Check Nginx is proxying correctly",
    "curl http://YOUR_VPS_IP/api/health",
    "",
    "# Tail live logs",
    "journalctl -u ibkr-backend -f",
    "tail -f ~/ibkr_trader/watchdog.log",
])

body("Open a browser and navigate to http://YOUR_VPS_IP — you should see the trading dashboard.")

# ═══════════════════════════════════════════════════════════════════════════════
# STEP 9 (OPTIONAL)
# ═══════════════════════════════════════════════════════════════════════════════
heading1("Step 9 (Optional) — SSL with Let's Encrypt")
body("If you have a domain name pointed at your VPS, add HTTPS with a free certificate:")

code_block([
    "sudo apt install certbot python3-certbot-nginx -y",
    "sudo certbot --nginx -d yourdomain.com",
    "# Auto-renewal is set up automatically via cron",
])

# ═══════════════════════════════════════════════════════════════════════════════
# USEFUL COMMANDS
# ═══════════════════════════════════════════════════════════════════════════════
heading1("Daily Operations Reference")

add_table(
    ["Task", "Command"],
    [
        ["View backend logs",           "journalctl -u ibkr-backend -f"],
        ["View gateway logs",           "journalctl -u ibgateway -f"],
        ["View watchdog log",           "tail -f ~/ibkr_trader/watchdog.log"],
        ["Restart backend",             "sudo systemctl restart ibkr-backend"],
        ["Restart everything",          "sudo systemctl restart ibgateway ibkr-backend ibkr-watchdog"],
        ["Check service status",        "sudo systemctl status ibkr-backend"],
        ["Pull code updates",           "cd ~/ibkr_trader && git pull"],
        ["After code update",           "sudo systemctl restart ibkr-backend"],
        ["Check IBKR connection",       "curl localhost:8000/status"],
        ["Manual retrain model",        "curl -X POST localhost:8000/journal/retrain"],
        ["Check trade journal stats",   "curl localhost:8000/journal/stats | python3 -m json.tool"],
        ["Switch to live account",      "Edit .env: TWS_PORT=4001, restart backend"],
    ]
)

# ═══════════════════════════════════════════════════════════════════════════════
# TROUBLESHOOTING
# ═══════════════════════════════════════════════════════════════════════════════
heading1("Troubleshooting")

heading2("Backend won't connect to IB Gateway")
bullet("Confirm IB Gateway is running: sudo systemctl status ibgateway")
bullet("Confirm port 4002 is listening: ss -tlnp | grep 4002")
bullet("Check IBC login succeeded: journalctl -u ibgateway -n 50")
bullet("Ensure 'Allow connections from localhost' is enabled in IB Gateway settings")

heading2("Frontend shows 'Not Connected'")
bullet("Check FastAPI is running: curl localhost:8000/health")
bullet("Check Nginx is proxying: curl http://YOUR_VPS_IP/api/health")
bullet("Check browser console for CORS or 404 errors")
bullet("Ensure BASE URL in index.html is set to '/api' not 'localhost:8000'")

heading2("IB Gateway asks for 2FA on startup")
bullet("Enable 'Trusted IPs' in IBKR account settings for your VPS IP")
bullet("Or use IBKR's SLS (Secure Login System) paper trading which skips 2FA")
bullet("IBC config: ReloginAfterSecondFactorAuthenticationTimeout=yes")

heading2("Watchdog keeps restarting the backend")
bullet("Backend is crashing — check: journalctl -u ibkr-backend -n 100")
bullet("Common cause: IB Gateway not ready yet when backend starts")
bullet("Fix: increase the 'sleep 30' to 'sleep 60' in the start sequence")

# ═══════════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════════
out = "C:/Projects/GenAI-Projects/ibkr_trader/ibkr_vps_deployment_guide.docx"
doc.save(out)
print(f"Saved: {out}")
