"""Generate IBKRAlgoTrader User Manual as a formatted .docx file."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ──────────────────────────────────────────────────────────────────
BRAND   = RGBColor(0x1A, 0x56, 0xDB)   # blue
DARK    = RGBColor(0x11, 0x18, 0x27)
GRAY    = RGBColor(0x6B, 0x72, 0x80)
GREEN   = RGBColor(0x05, 0x7A, 0x55)
HDR_BG  = "1A56DB"                      # table header fill
ALT_BG  = "EBF5FF"                      # alternating row fill

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, attrs in kwargs.items():
        tag = OxmlElement(f"w:{ edge }")
        for k, v in attrs.items():
            tag.set(qn(f"w:{k}"), v)
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(hdr_cells[i], HDR_BG)

    # data rows
    for ri, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        bg = ALT_BG if ri % 2 == 1 else "FFFFFF"
        for ci, val in enumerate(row_data):
            row_cells[ci].text = val
            row_cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
            set_cell_bg(row_cells[ci], bg)

    # col widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size  = Pt(22)
    run.font.bold  = True
    run.font.color.rgb = BRAND
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.size  = Pt(14)
    run.font.bold  = True
    run.font.color.rgb = DARK
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size  = Pt(11)
    run.font.bold  = True
    run.font.color.rgb = GRAY
    return p

def body(text, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size   = Pt(10)
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run("ℹ  " + text)
    run.font.size   = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x09, 0x53, 0x90)
    return p

def code_block(text):
    for line in text.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.3)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.size = Pt(10)
        run.font.bold = True
        run2 = p.add_run(text)
        run2.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)

def divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)

# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run("IBKR AlgoTrader")
run.font.size  = Pt(32)
run.font.bold  = True
run.font.color.rgb = BRAND

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("User Manual — Version 1.0")
run2.font.size  = Pt(14)
run2.font.color.rgb = GRAY

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("Automated Options Trading for Interactive Brokers")
run3.font.size   = Pt(11)
run3.font.italic = True
run3.font.color.rgb = GRAY

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (static)
# ════════════════════════════════════════════════════════════════════════════
h1("Table of Contents")
toc = [
    ("1", "Prerequisites"),
    ("2", "Download & Install"),
    ("3", "License Activation"),
    ("4", "TWS Configuration (IBKR Side)"),
    ("5", "App Connection Setup"),
    ("6", "Paper Trading — First Run"),
    ("7", "SPX 0DTE Strategy Guide"),
    ("8", "Troubleshooting"),
    ("9", "Upgrading"),
]
for num, title in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"  {num}.  {title}")
    run.font.size = Pt(10)
    run.font.color.rgb = BRAND

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 1. PREREQUISITES
# ════════════════════════════════════════════════════════════════════════════
h1("1.  Prerequisites")
body("Before installing, make sure you have the following in place:")
add_table(
    ["Requirement", "Details"],
    [
        ["IBKR Account",               "Paper or Live. Open at ibkr.com"],
        ["Trader Workstation (TWS)",   "Download from ibkr.com/trading/tws — Version 10.19 or later"],
        ["Operating System",           "Windows 10 or 11, 64-bit"],
        ["Options Trading Enabled",    "Must be approved for options on your IBKR account"],
        ["License Key",                "Received via email after purchase at [your product URL]"],
    ],
    col_widths=[2.0, 4.2],
)
note("Strongly recommended: complete at least 2 weeks of paper trading before switching to a live account.")

divider()

# ════════════════════════════════════════════════════════════════════════════
# 2. DOWNLOAD & INSTALL
# ════════════════════════════════════════════════════════════════════════════
h1("2.  Download & Install")

h3("Step 1 — Download")
body("Go to your product website and download the latest installer:")
code_block("IBKRAlgoTrader-v1.0-Setup.exe")

h3("Step 2 — Run the Installer")
bullet("Double-click the installer file.")
bullet("Accept the license agreement.")
bullet("Choose install location (default: C:\\Program Files\\IBKRAlgoTrader\\).")
bullet("Click Install and wait for completion.")
bullet('When finished, check "Launch IBKRAlgoTrader" and click Finish.')

h3("Step 3 — First Launch")
body("The app starts a local server on your machine. Your default browser will open automatically to:")
code_block("http://localhost:8000")
body("Keep this browser tab open while trading. The app continues running in the background — look for the IBKRAlgoTrader icon in your Windows system tray (bottom-right corner of your taskbar).")

divider()

# ════════════════════════════════════════════════════════════════════════════
# 3. LICENSE ACTIVATION
# ════════════════════════════════════════════════════════════════════════════
h1("3.  License Activation")

h3("Step 1 — Enter Your License Key")
body("On first launch you will see the Activation screen. Paste the license key from your purchase confirmation email:")
code_block("ALGO-XXXX-XXXX-XXXX-XXXX")

h3("Step 2 — Activate")
body("Click Activate. The app contacts the license server to verify your key and unlock your features.")

h3("Step 3 — Feature Unlock")
body("What is available depends on your plan:")
add_table(
    ["Plan", "Features Included"],
    [
        ["Starter", "Portfolio Dashboard + SPX 0DTE Auto-Trader"],
        ["Pro",     "Starter + Earnings Volatility Crush (EVC)"],
        ["Elite",   "Pro + Breakout Scanner + Day Trader"],
    ],
    col_widths=[1.5, 4.7],
)

h3("Step 4 — Confirmation")
body('A green "License Active" banner appears at the top of the app showing your plan name and renewal date. Only the tabs included in your plan are visible in the navigation.')
note("The app re-validates your license on every startup. An internet connection is required at launch.")

divider()

# ════════════════════════════════════════════════════════════════════════════
# 4. TWS CONFIGURATION
# ════════════════════════════════════════════════════════════════════════════
h1("4.  TWS Configuration (IBKR Side)")
body("This is the most critical setup step. You must enable API access inside Trader Workstation before the app can connect and place orders.")

h3("Step 1 — Open TWS and Log In")
body("Launch Trader Workstation and sign in with your IBKR credentials. For paper trading, use your paper account credentials (separate login from live).")

h3("Step 2 — Open API Settings")
body("In the TWS menu bar, navigate to:")
code_block("Edit  →  Global Configuration  →  API  →  Settings")

h3("Step 3 — Apply These Settings")
add_table(
    ["Setting", "Value", "Notes"],
    [
        ["Enable ActiveX and Socket Clients", "✅ Checked",   "Required — enables the API connection"],
        ["Socket port",                        "7497 / 7496", "7497 = paper trading  |  7496 = live trading"],
        ["Allow connections from localhost",   "✅ Checked",   "Security — only your machine can connect"],
        ["Read-Only API",                      "❌ Unchecked", "Must be OFF so the app can place orders"],
        ["Master API client ID",               "Leave blank",  "No change needed"],
        ["Trusted IP Addresses",               "127.0.0.1",   "Add this entry"],
    ],
    col_widths=[2.5, 1.4, 2.3],
)

h3("Step 4 — Save and Restart TWS")
body("Click OK to save, then fully close and reopen TWS. TWS must be running and logged in before you start the AlgoTrader app each day.")

note("If you see a pop-up in TWS saying 'Do you want to allow API connections from...', click Yes.")

divider()

# ════════════════════════════════════════════════════════════════════════════
# 5. APP CONNECTION SETUP
# ════════════════════════════════════════════════════════════════════════════
h1("5.  App Connection Setup")

h3("Step 1 — Open Connection Settings")
body("In the app, click the gear icon (top right) → Connection.")

h3("Step 2 — Enter TWS Details")
add_table(
    ["Field", "Paper Trading", "Live Trading"],
    [
        ["TWS Host",  "127.0.0.1", "127.0.0.1"],
        ["TWS Port",  "7497",      "7496"],
        ["Client ID", "1",         "1"],
    ],
    col_widths=[1.8, 2.1, 2.1],
)

h3("Step 3 — Connect")
body("Click Connect to TWS.")
bullet("Success: The top banner turns green and shows your account number and net liquidation value.")
bullet("Failure: See Section 8 — Troubleshooting.")

h3("Step 4 — Verify Account Data")
body("Navigate to the Portfolio tab. Confirm that your positions and buying power match what you see in TWS. If they match, the connection is working correctly.")

divider()

# ════════════════════════════════════════════════════════════════════════════
# 6. PAPER TRADING FIRST RUN
# ════════════════════════════════════════════════════════════════════════════
h1("6.  Paper Trading — First Run")
body("Before enabling any live strategy, complete these verification steps:")

bullet("Confirm you are on the paper port (7497) — the connection banner shows (PAPER) next to the account number.", None)
bullet("Enable the strategy you want to test but observe it for at least 5 market days before going live.", None)
bullet("Check the Trade Journal after each day — verify fills, P&L, and exit timing look correct.", None)
bullet("Cross-check TWS Order History — orders in TWS should match what the app logged.", None)
bullet("Only switch to live port (7496) after you are satisfied with paper performance.", None)

note("Paper trading uses simulated fills. Actual live fills may differ slightly in fast-moving markets.")

divider()

# ════════════════════════════════════════════════════════════════════════════
# 7. SPX 0DTE STRATEGY GUIDE
# ════════════════════════════════════════════════════════════════════════════
h1("7.  SPX 0DTE Strategy Guide")

h3("What It Does")
body("Automatically sells an SPX iron condor each morning and manages it through the trading day. Tiered profit targets ($200 → $150 → $100 → $50) allow up to 4 separate trades per day, each sized to hit its own target independently.")

h3("Enable the Strategy")
bullet("Click the SPX 0DTE tab in the navigation bar.")
bullet("Review the default parameters (see table below).")
bullet("Toggle Enable to ON.")

h3("Key Parameters")
add_table(
    ["Parameter", "Default", "Description"],
    [
        ["Max Margin",     "$5,000",            "Maximum capital allocated per spread"],
        ["Profit %",       "50%",               "Close spread when this % of premium is captured"],
        ["Trade Targets",  "$200 / $150 / $100 / $50", "P&L goal for each successive attempt"],
        ["Max Attempts",   "4",                 "Maximum iron condor entries per day"],
        ["Entry Window",   "9:31–9:45 AM ET",   "Time window when first condor is placed"],
        ["Delta Target",   "0.10",              "Target delta for short legs at entry"],
    ],
    col_widths=[1.8, 1.8, 2.6],
)

h3("Daily Automation Schedule")
add_table(
    ["Time (ET)", "Action"],
    [
        ["9:31 AM",       "App scans SPX chain and places first iron condor"],
        ["Throughout day","Monitors P&L every 30 seconds"],
        ["On profit hit", "Closes spread, logs to journal, re-enters for next target if attempts remain"],
        ["3:45 PM",       "Force-closes any remaining open spreads before market close"],
    ],
    col_widths=[1.5, 4.7],
)

h3("Monitoring")
body("The Open Iron Condors table shows live P&L, strike prices, and collected credit for each active spread. The Decisions Log below it records every automated action with timestamps — useful for reviewing strategy behavior.")

h3("Trade Journal")
body("Go to SPX 0DTE → History to see all closed trades: entry price, exit price, P&L, win/loss, and exit reason (profit target, max loss, or end-of-day close).")

divider()

# ════════════════════════════════════════════════════════════════════════════
# 8. TROUBLESHOOTING
# ════════════════════════════════════════════════════════════════════════════
h1("8.  Troubleshooting")

h2('"Cannot connect to TWS"')
bullet("Confirm TWS is open and you are logged in.")
bullet("Confirm the port matches: 7497 for paper, 7496 for live.")
bullet("In TWS: Edit → Global Configuration → API → Settings → verify Enable ActiveX and Socket Clients is checked.")
bullet("Restart both TWS and the AlgoTrader app, then try connecting again.")

h2('"License invalid" or "Cannot reach license server"')
bullet("Check your internet connection — the app requires internet access at startup.")
bullet("Verify the license key was pasted without extra spaces.")
bullet("If your subscription has lapsed, renew at [your product URL].")
bullet("Contact support: support@[yourproduct].com")

h2("App opens but browser shows blank / cannot connect")
bullet("Manually open your browser and go to http://localhost:8000")
bullet("Check the system tray for the IBKRAlgoTrader icon — right-click → Open Dashboard.")
bullet("If missing from tray, restart the app from the Start Menu.")

h2("Spreads not filling")
bullet("Confirm your IBKR account has options trading permissions enabled.")
bullet("Check buying power — insufficient margin will cause orders to be rejected by IBKR.")
bullet("In TWS, check for pending order confirmation dialogs (disable these in TWS API settings).")

h2("SPX price looks different from other platforms")
body("SPX live price is derived from SPY using a ratio to the previous close. A brief delay in the IBKR market data feed can cause a momentary mismatch — this resolves within a few seconds and does not affect order pricing.")

divider()

# ════════════════════════════════════════════════════════════════════════════
# 9. UPGRADING
# ════════════════════════════════════════════════════════════════════════════
h1("9.  Upgrading")

h3("Automatic Update Notification")
body("On startup, the app checks for a newer version. If one is available, a banner appears at the top:")
code_block("⬆  Version 1.1 available — Download Now")
body("Click Download Now. The new installer downloads and launches automatically. Your license key, connection settings, and trade history are all preserved.")

h3("Manual Upgrade")
bullet("Download the latest installer from [your product URL].")
bullet("Run the installer — it upgrades in place, no uninstall required.")
bullet("Your license key, TWS settings, and trade journal carry over automatically.")

h3("New Features After Upgrade")
body("After upgrading, any new features included in your plan appear automatically as new tabs in the navigation — no reconfiguration needed. Features not yet in your plan remain hidden until your plan is upgraded.")

note("You can check your current app version at any time: gear icon → About.")

doc.add_page_break()

# ── Footer ───────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("IBKR AlgoTrader v1.0  ·  support@[yourproduct].com  ·  docs.[yourproduct].com")
run.font.size  = Pt(8)
run.font.color.rgb = GRAY

# ── Save ─────────────────────────────────────────────────────────────────────
out_path = os.path.join(os.path.dirname(__file__), "IBKRAlgoTrader_UserManual_v1.0.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
